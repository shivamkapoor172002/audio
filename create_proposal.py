from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

images_dir = r'c:\Users\shiva\OneDrive\Desktop\audionix\static\images'

def add_image_safe(doc, img_name, width=5):
    """Safely add an image, skip if it fails"""
    img_path = os.path.join(images_dir, img_name)
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(width))
            return True
        except Exception as e:
            doc.add_paragraph(f'[Image: {img_name}]')
            return False
    return False

# ========== COVER PAGE ==========
add_image_safe(doc, 'audionix_Main_logo-removebg-preview.png', 2.5)
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER if doc.paragraphs else None

doc.add_paragraph()

title = doc.add_heading('AUDIONIX ENGINEERS LLP', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Experiences, Amplified.')
run.font.size = Pt(18)
run.font.italic = True
run.font.color.rgb = RGBColor(201, 169, 98)

doc.add_paragraph()

doc_title = doc.add_paragraph()
doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doc_title.add_run('Website Content & Design Overview')
run.font.size = Pt(20)
run.font.bold = True

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run('February 2026')

doc.add_page_break()

# ========== COLOR THEME ==========
doc.add_heading('Color Theme & Design', level=1)
doc.add_paragraph('The website uses an elegant Cream & Navy Blue color palette:')

colors_data = [
    ('Primary Background', '#E8E0D0', 'Warm cream - main background'),
    ('Navy Primary', '#1A2B5A', 'Deep navy - headers, buttons'),
    ('Navy Dark', '#0A1628', 'Darker navy - footer'),
    ('Gold Accent', '#C9A962', 'Gold - highlights, decorations'),
    ('Cream Light', '#F5F0E6', 'Light cream - cards'),
]

color_table = doc.add_table(rows=1, cols=3)
color_table.style = 'Table Grid'
hdr = color_table.rows[0].cells
hdr[0].text = 'Color'
hdr[1].text = 'Hex Code'
hdr[2].text = 'Usage'

for name, hex_code, usage in colors_data:
    row = color_table.add_row().cells
    row[0].text = name
    row[1].text = hex_code
    row[2].text = usage

doc.add_paragraph()
doc.add_heading('Typography', level=2)
doc.add_paragraph('• Playfair Display - Headings and titles')
doc.add_paragraph('• Montserrat - Body text and navigation')
doc.add_paragraph('• Great Vibes - Decorative script taglines')

doc.add_page_break()

# ========== HOME PAGE ==========
doc.add_heading('HOME PAGE', level=1)

doc.add_heading('Hero Section', level=2)
add_image_safe(doc, 'speakers_image.jpg', 5)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Welcome Message: ').bold = True
p.add_run('"Welcome to Audionix Engineering LLP"')

p = doc.add_paragraph()
p.add_run('Tagline: ').bold = True
p.add_run('"Experiences, Amplified."').italic = True

doc.add_paragraph()
doc.add_heading('Key Statistics:', level=3)
doc.add_paragraph('• 10+ Years Experience')
doc.add_paragraph('• 150+ Projects Completed')
doc.add_paragraph('• 30+ Cities Served')

doc.add_paragraph()
doc.add_heading('Who We Are Section', level=2)

add_image_safe(doc, 'speakers_with_screen.jpg', 4)
doc.add_paragraph()
add_image_safe(doc, 'Auditorium.jpg', 4)

doc.add_paragraph()
doc.add_paragraph(
    'AUDIONIX ENGINEERS LLP is an independent consultancy that designs professional AV, acoustics and lighting systems '
    'for government, institutional and corporate spaces across India. With deep on-ground experience and a clear, '
    'vendor-neutral approach, we bridge the gap between client expectations, real-world site challenges and rapidly evolving technology.'
)

p = doc.add_paragraph()
p.add_run('The Result? ').bold = True
p.add_run('Solutions that are reliable, scalable and effortless to operate.')

doc.add_page_break()

# About Founder Section
doc.add_heading('About Founder Section', level=2)
add_image_safe(doc, 'founder_owner_image_vineet.jpg', 2.5)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Vineet Tripathi').bold = True
p.add_run(' - Founder')

doc.add_paragraph(
    'With over 10+ years of experience in India\'s professional AV industry, Vineet Tripathi brings deep technical '
    'expertise and hands-on experience from working with global leaders like Bose Professional and Harman International.'
)

doc.add_paragraph(
    'His career spans complex, large-format projects across government, institutional and corporate environments, '
    'giving him a clear understanding of real-site challenges, tendering processes and performance expectations.'
)

doc.add_page_break()

# ========== SERVICES ==========
doc.add_heading('SERVICES PAGE', level=1)

p = doc.add_paragraph()
p.add_run('Page Tagline: ').bold = True
p.add_run('"Where sound, sight and sense come together."').italic = True

doc.add_paragraph()

# Service 1
doc.add_heading('1. AV System Design & Consulting', level=2)
add_image_safe(doc, 'speakers_with_screen.jpg', 4)
doc.add_paragraph(
    'End-to-end design for auditoriums, lecture halls, boardrooms, convention centres and public spaces. '
    'From requirement mapping and schematics to BoQs, drawings and tender-ready documents.'
)

# Service 2
doc.add_heading('2. Room Acoustics', level=2)
add_image_safe(doc, 'Auditorium.jpg', 4)
doc.add_paragraph(
    'End-to-end acoustic planning including analysis, material selection and layout optimisation to deliver '
    'crisp speech clarity, balanced sound and an environment tailored to its purpose.'
)

# Service 3
doc.add_heading('3. Architectural & Stage Lighting Design', level=2)
add_image_safe(doc, 'a_image_stage_spotlight.jpg', 4)
doc.add_paragraph(
    'Thoughtfully crafted lighting solutions that elevate both function and ambience. Ideal for auditoriums, '
    'multipurpose halls and performance spaces, with a focus on comfort, visibility and energy efficiency.'
)

doc.add_page_break()

# Service 4
doc.add_heading('4. Technical Review & Value Engineering', level=2)
doc.add_paragraph(
    'Independent audits of OEM or contractor proposals to ensure accuracy, fairness and performance. '
    'We identify gaps, remove unnecessary costs and strengthen the overall system design.'
)

# Service 5
doc.add_heading('5. Information & Communication Technology', level=2)
add_image_safe(doc, 'rbi_office.jpg', 4)
doc.add_paragraph(
    'Planning and design of IT and networking infrastructure for new and existing buildings, ensuring '
    'seamless connectivity, efficient communication and future-ready systems.'
)

doc.add_page_break()

# ========== PORTFOLIO ==========
doc.add_heading('PORTFOLIO PAGE', level=1)

p = doc.add_paragraph()
p.add_run('Title: ').bold = True
p.add_run('The Leadership Portfolio')

p = doc.add_paragraph()
p.add_run('Subtitle: ').bold = True
p.add_run('"Led With Precision"').italic = True

doc.add_paragraph(
    'From high-stakes courtrooms and national institutions to smart cities and sacred spaces, '
    'Vineet Tripathi\'s portfolio spans some of India\'s most demanding environments.'
)

doc.add_heading('Flagship Projects', level=2)

# RBI
doc.add_heading('RBI (Reserve Bank of India)', level=3)
add_image_safe(doc, 'rbi_office.jpg', 4)
p = doc.add_paragraph()
p.add_run('Pan-India Boardroom Modernisation').bold = True
doc.add_paragraph('Scale: 150+ Boardrooms across 30+ Cities')
doc.add_paragraph('One of the biggest AMX deployments nationwide.')

# Prayagraj
doc.add_heading('Prayagraj Smart City', level=3)
add_image_safe(doc, 'Indira_paryavaran_bhawan.jpg', 4)
p = doc.add_paragraph()
p.add_run('Kumbh Mela 2025 Infrastructure').bold = True
doc.add_paragraph('Scale: 700+ Speakers, 200+ Amplifiers, 200+ Intersections')

doc.add_page_break()

doc.add_heading('Sector Experience', level=2)

sectors = [
    ('Convention Centres', 'Bharat_mandapam.jpg', 'ITPO Delhi, Minto Hall Bhopal, ICAR Convention Centre'),
    ('Public Infrastructure', 'Indira_paryavaran_bhawan.jpg', 'National War Memorial, National Police Memorial, Vigyan Bhawan'),
    ('Judiciary', 'aiims_rae_bareli.jpg', 'Supreme Court Extension, District Court Goa'),
    ('Places of Worship', 'Kushabhau_thakre_hall.jpg', 'Ganga Ghats Varanasi, Kedarnath Temple, Kashi Vishwanath Corridor'),
    ('BFSI', 'rbi_office.jpg', 'RBI, SBI'),
    ('Experience Centres', 'Banjara_virsat_museum.jpg', 'Banjara Museum, Science City Ahmedabad, Atal Sarovar'),
    ('Education', 'aiims_rae_bareli.jpg', 'IIT Bhilai, Jodhpur, Roorkee | AIIMS Jodhpur | IIM Jammu'),
]

for sector_name, img_file, projects in sectors:
    p = doc.add_paragraph()
    p.add_run(f'• {sector_name}: ').bold = True
    p.add_run(projects)

doc.add_page_break()

# ========== ABOUT PAGE ==========
doc.add_heading('ABOUT & FOUNDER PAGE', level=1)

add_image_safe(doc, 'auditorium_view.jpg', 5)
doc.add_paragraph()
add_image_safe(doc, 'Auditorium_view_eagel_eye.jpg', 5)

doc.add_paragraph()
doc.add_heading('The Audionix Advantage', level=2)

advantages = [
    ('Seamless Collaboration', 'We work smoothly with architects, PMCs, MEP, IT, and contractors.'),
    ('End-to-End Ownership', 'From requirement analysis to BoQs, drawings, tuning, and site reviews.'),
    ('Needs-First Thinking', 'We begin with your functional, architectural, and operational requirements.'),
    ('Product-Agnostic Expertise', 'Across major OEMs, we pick what fits best.')
]

for title, desc in advantages:
    p = doc.add_paragraph()
    p.add_run(f'✓ {title}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== CONTACT PAGE ==========
doc.add_heading('CONTACT PAGE', level=1)

p = doc.add_paragraph()
p.add_run('Tagline: ').bold = True
p.add_run('"Let\'s engineer something exceptional together."').italic = True

doc.add_paragraph()
doc.add_heading('Contact Details', level=2)
doc.add_paragraph('📧 Email: INFO@AUDIONIXENGINEERS.COM')
doc.add_paragraph('📞 Phone: +91 99997 62624')
doc.add_paragraph('📍 Location: New Delhi, India')
doc.add_paragraph('🌐 Serving clients pan-India')

doc.add_paragraph()
doc.add_heading('Contact Form', level=2)
doc.add_paragraph('Fields: Name, Email, Subject, Message')
doc.add_paragraph('Includes embedded Google Maps for location.')

doc.add_page_break()

# ========== SUMMARY ==========
doc.add_heading('Summary', level=1)

summary = doc.add_table(rows=6, cols=2)
summary.style = 'Table Grid'
data = [
    ('Total Pages', '5 - Home, About, Services, Portfolio, Contact'),
    ('Design Style', 'Professional cream & navy theme'),
    ('Target Clients', 'Government, Institutional, Corporate'),
    ('Key Features', 'Responsive, animations, contact form'),
    ('Technology', 'Python Flask, HTML5, CSS3, JavaScript'),
    ('Contact', 'info@audionixengineers.com | +91 99997 62624')
]
for i, (key, val) in enumerate(data):
    summary.rows[i].cells[0].text = key
    summary.rows[i].cells[1].text = val

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('AUDIONIX ENGINEERS LLP').bold = True
footer.add_run(' | Experiences, Amplified.')

# Save
import datetime
timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
output_path = rf'c:\Users\shiva\OneDrive\Desktop\audionix\Audionix_Website_Proposal_{timestamp}.docx'

try:
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
except PermissionError:
    print(f"Error: Could not save to {output_path}. Please close the file if it is open.")
    # Try a fallback
    fallback_path = rf'c:\Users\shiva\OneDrive\Desktop\audionix\Audionix_Proposal_New.docx'
    doc.save(fallback_path)
    print(f'Document saved to: {fallback_path}')

